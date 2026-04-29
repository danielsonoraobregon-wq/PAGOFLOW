import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from num2words import num2words
    def _n2w(n, **kw):
        return num2words(int(n), lang='es', **kw)
except ImportError:
    def _n2w(n, **kw):
        return str(n)

MESES = ['enero','febrero','marzo','abril','mayo','junio',
         'julio','agosto','septiembre','octubre','noviembre','diciembre']


def monto_letras(n: float) -> str:
    entero = int(n)
    centavos = round((n - entero) * 100)
    texto = _n2w(entero).upper()
    if centavos:
        return f"{texto} PESOS {centavos:02d}/100 M.N."
    return f"{texto} PESOS 00/100 M.N."


def ordinal_es(n: int) -> str:
    return _n2w(n, to='ordinal').upper()


def generar_recibo(
    lote: dict,
    proyecto: dict,
    contrato: dict,
    num_mensualidad: int,
    monto: float,
    vendedor_nombre: str,
    ciudad: str,
    estado: str,
) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    total_mens  = int(contrato.get('total_mensualidades') or 0)
    monto_total = float(contrato.get('monto_total') or 0)
    superficie  = contrato.get('superficie') or ''
    ubicacion   = contrato.get('ubicacion') or proyecto.get('descripcion') or ''

    now       = datetime.now()
    fecha_str = f"{now.day} de {MESES[now.month - 1]} del año {now.year}"
    ordinal   = ordinal_es(num_mensualidad)

    # Título
    titulo = (
        f"RECIBO DE PAGO LOTE N.º {lote['numero']} "
        f"{proyecto['nombre'].upper()} "
        f"{ordinal} MENSUALIDAD."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Fecha
    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fecha.add_run(f"{ciudad}, {estado} a {fecha_str}.")

    doc.add_paragraph()

    # Cuerpo
    cuerpo = (
        f"Recibí del C. {lote['cliente_nombre']} la cantidad de "
        f"${monto:,.0f} ({monto_letras(monto)}), "
        f"por concepto de la {ordinal} mensualidad de un total de "
        f"{total_mens} mensualidades por concepto compra venta del terreno "
        f"número {lote['numero']} de mi propiedad ubicado en el proyecto "
        f"{proyecto['nombre']}"
    )
    if ubicacion:
        cuerpo += f" en el área de {ubicacion}"
    cuerpo += (
        f", marcado con el número ({lote['numero']}) del citado proyecto "
        f"{proyecto['nombre']}"
    )
    if superficie:
        cuerpo += f" con una Superficie Total de {superficie}"
    cuerpo += (
        f", dicha cantidad de ${monto:,.0f} ({monto_letras(monto)}), "
        f"es un abono del inmueble que tiene un costo total de "
        f"${monto_total:,.0f} ({monto_letras(monto_total)})."
    )

    p_body = doc.add_paragraph(cuerpo)
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for _ in range(4):
        doc.add_paragraph()

    # Firma
    for txt, bold in [
        ("_______________________________", False),
        ("Recibe de conformidad",           False),
        (vendedor_nombre,                   True),
    ]:
        p_f = doc.add_paragraph(txt)
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bold:
            p_f.runs[0].bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
